from pathlib import Path
from PySide6.QtWidgets import(QApplication, QFileDialog, QMainWindow, QVBoxLayout,
                               QWidget, QTextEdit, QListWidget, QSplitter, QLabel,
                               QPushButton,QAbstractItemView,QTableWidgetItem,QTableWidget, QMessageBox, QHBoxLayout, QComboBox,QListWidgetItem)
from PySide6.QtUiTools import QUiLoader
from PySide6.QtCore import Qt,QEvent, QObject
import sys
import shutil
import re
from docx import Document
import pdfplumber
import pandas

class DragnDrop(QObject):
    def __init__(self, window):
        super().__init__()
        self.window = window
        self.files_folder = Path("input_files")
        self.templates_folder = Path("templates")
        self.current_file_path = None
        self.current_template_path = None
        self.setup()

    def setup(self):
        self.files_folder.mkdir(exist_ok=True)
        self.templates_folder.mkdir(exist_ok=True)
        if hasattr(self.window, 'leftDropZone'):
            self.window.leftDropZone.setAcceptDrops(True)
            self.window.leftDropZone.installEventFilter(self)
        if hasattr(self.window, 'rightDropZone'):
            self.window.rightDropZone.setAcceptDrops(True)
            self.window.rightDropZone.installEventFilter(self)
        if hasattr(self.window, 'File'):
            self.window.File.triggered.connect(self.load_file_from_menu)
        if hasattr(self.window, 'Template'):
            self.window.Template.triggered.connect(self.load_template_from_menu)

    def eventFilter(self, obj, event):
        if event.type() == QEvent.DragEnter:
            if obj == self.window.leftDropZone:
                self.drag_enter_files(event)
            elif obj == self.window.rightDropZone:
                self.drag_enter_templates(event)
            return True
        elif event.type() == QEvent.Drop:
            if obj == self.window.leftDropZone:
                self.drop_files(event)
            elif obj == self.window.rightDropZone:
                self.drop_templates(event)
            return True
        return super().eventFilter(obj, event)

    def drag_enter_files(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def drop_files(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                source = Path(url.toLocalFile())
                if source.is_file():
                    dest = self.copy_file_to_folder(source, self.files_folder)
                    self.current_file_path = dest
                    self.show_content_in_edit(dest, self.window.leftContentEdit)
            self.window.statusbar.showMessage("Файлы добавлены в input_files", 2000)
        event.acceptProposedAction()

    def drop_templates(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                source = Path(url.toLocalFile())
                if source.is_file():
                    dest = self.copy_file_to_folder(source, self.templates_folder)
                    self.current_template_path = dest
                    self.show_content_in_edit(dest, self.window.rightContentEdit)
            self.window.statusbar.showMessage("Шаблоны добавлены в templates", 2000)
        event.acceptProposedAction()

    def show_content_in_edit(self, file_path: Path, text_edit):
        suffix = file_path.suffix.lower()
        content = ""
        try:
            if suffix == '.docx':
                doc = Document(file_path)
                paragraphs = [para.text for para in doc.paragraphs]
                content = "\n".join(paragraphs)
            elif suffix == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    content = text
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='cp1251') as f:
                    content = f.read()
            except Exception as e:
                content = f"Не удалось прочитать файл как текст: {e}"
        except Exception as e:
            content = f"Не удалось прочитать файл:\n{e}"
        text_edit.setPlainText(content)

    def copy_file_to_folder(self, source_path: Path, target_folder: Path) -> Path:
        source_path = source_path.resolve()
        target_folder = target_folder.resolve()
        if source_path.parent == target_folder:
            return source_path
        dest_path = target_folder / source_path.name
        if dest_path.exists():
            stem = source_path.stem
            suffix = source_path.suffix
            counter = 1
            while dest_path.exists():
                dest_path = target_folder / f"{stem}_{counter}{suffix}"
                counter += 1
        shutil.copy2(source_path, dest_path)
        return dest_path

    def load_file_from_menu(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self.window,
            "Выберите файл",
            str(self.files_folder),
            "Все файлы (*.*)"
        )
        if file_path:
            dest_path = self.copy_file_to_folder(Path(file_path), self.files_folder)
            self.current_file_path = dest_path
            self.show_content_in_edit(dest_path, self.window.leftContentEdit)
            self.window.statusbar.showMessage(f"Файл скопирован в {self.files_folder}", 2000)

    def load_template_from_menu(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self.window,
            "Выберите шаблон",
            str(self.templates_folder),
            "Все файлы (*.*)"
        )
        if file_path:
            dest_path = self.copy_file_to_folder(Path(file_path), self.templates_folder)
            self.current_template_path = dest_path
            self.show_content_in_edit(dest_path, self.window.rightContentEdit)
            self.window.statusbar.showMessage(f"Шаблон скопирован в {self.templates_folder}", 2000)

class FillWindowTemplate(QMainWindow):
    def __init__(self, path, dataframe=None, column_mapping=None, parent=None):
        super().__init__(parent)
        self.path = path
        self.dataframe = dataframe
        self.column_mapping = column_mapping or {}
        self.markers = []
        self.variant_widgets = []
        self.generated_text = ""
        self.setWindowTitle(f"Заполнение шаблона: {path.name}")
        self.resize(1000, 700)
        self.setup_ui()
        self.load_template()

    def setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)

        top_layout = QHBoxLayout()
        self.example_btn = QPushButton("Пример")
        self.generate_btn = QPushButton("Сгенерировать")
        top_layout.addWidget(self.example_btn)
        top_layout.addWidget(self.generate_btn)

        top_layout.addStretch()
        top_layout.addWidget(QLabel("Имя файла:"))
        self.file_name_combo = QComboBox()
        self.file_name_combo.addItem("Имя шаблона")
        top_layout.addWidget(self.file_name_combo)

        layout.addLayout(top_layout)

        splitter = QSplitter(Qt.Horizontal)
        layout.addWidget(splitter)

        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(False)
        self.text_edit.textChanged.connect(self.on_text_changed)
        splitter.addWidget(self.text_edit)

        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("Маркеры и теги:"))
        self.markers_list = QListWidget()
        right_layout.addWidget(self.markers_list)

        del_btn = QPushButton("Удалить маркер")
        del_btn.clicked.connect(self.remove_marker)
        right_layout.addWidget(del_btn)

        splitter.addWidget(right_widget)
        splitter.setSizes([600, 300])

        self.example_btn.clicked.connect(self.show_example)
        self.generate_btn.clicked.connect(self.generate_all)

    def load_template(self):
        suffix = self.path.suffix.lower()
        if suffix == '.docx':
            self.load_docx()
        elif suffix == '.pdf':
            self.load_pdf()
        else:
            self.load_text()
        self.display_text_with_markers()
        self.populate_variants_panel()
        self.update_file_name_options()

    def load_docx(self):
        doc = Document(self.path)
        full_text = [para.text for para in doc.paragraphs]
        self.raw_text = "\n".join(full_text)
        self.find_markers_in_text(self.raw_text)

    def load_pdf(self):
        with pdfplumber.open(self.path) as pdf:
            full_text = [page.extract_text() or "" for page in pdf.pages]
        self.raw_text = "\n".join(full_text)
        self.find_markers_in_text(self.raw_text)

    def load_text(self):
        try:
            with open(self.path, 'r', encoding='utf-8') as f:
                self.raw_text = f.read()
        except Exception as e:
            self.raw_text = f"Ошибка чтения: {e}"
        self.find_markers_in_text(self.raw_text)

    def find_markers_in_text(self, text):
        underscore_pattern = r'_{2,}'
        bracket_pattern = r'<([^>]+)>'
        colon_pattern = r'([A-Za-zА-Яа-я0-9_$]+)\s*:'
        markers = []
        for match in re.finditer(underscore_pattern, text):
            start, end = match.span()
            word = self.find_word(text, start)
            markers.append({
                'type': 'underscore',
                'start': start,
                'end': end,
                'text': match.group(),
                'word': word
            })
        for match in re.finditer(bracket_pattern, text):
            start, end = match.span()
            inner = match.group(1)
            markers.append({
                'type': 'bracket',
                'start': start,
                'end': end,
                'text': match.group(),
                'word': inner.strip()
            })
        existing_starts = {m['start'] for m in markers}
        existing_ends = {m['end'] for m in markers}
        for match in re.finditer(colon_pattern, text):
            start, end = match.span()
            word = match.group(1)
            remainder = text[end:]
            next_newline = remainder.find('\n')
            rest_of_line = remainder[:next_newline] if next_newline != -1 else remainder
            if rest_of_line:
                if not re.match(r'^[ \t]*$', rest_of_line):
                    continue
            else:
                pass
            overlapped = False
            for s, e in zip(existing_starts, existing_ends):
                if not (end <= s or start >= e):
                    overlapped = True
                    break
            if overlapped:
                continue
            markers.append({
                'type': 'colon',
                'start': start,
                'end': end,
                'text': match.group(0),
                'word': word
            })
        markers.sort(key=lambda x: x['start'])
        self.markers = markers

    def find_word(self, text, pos):
        i = pos - 1
        while i >= 0 and (text[i].isspace() or text[i] in '.,:;!?()[]{}-'):
            i -= 1
        end = i
        while i >= 0 and (text[i].isalnum() or text[i] == '_'):
            i -= 1
        start = i + 1
        if start <= end:
            return text[start:end + 1]
        return ""

    def on_text_changed(self):
        self.raw_text = self.text_edit.toPlainText()
        self.find_markers_in_text(self.raw_text)
        self.populate_variants_panel()
        self.markers_list.repaint()

    def remove_marker(self):
        current_row = self.markers_list.currentRow()
        if current_row < 0:
            QMessageBox.information(self, "Информация", "Выберите маркер для удаления.")
            return
        del self.markers[current_row]
        self.populate_variants_panel()

    def display_text_with_markers(self):
        self.text_edit.setPlainText(self.raw_text)

    def populate_variants_panel(self):
        self.markers_list.clear()
        self.variant_widgets.clear()
        if self.dataframe is not None:
            columns = list(self.column_mapping.values())
        else:
            columns = []
        for marker in self.markers:
            widget = QWidget()
            layout = QHBoxLayout(widget)
            label = QLabel(marker['word'])
            combo = QComboBox()
            if columns:
                combo.addItems(columns)
                combo.setCurrentIndex(0)
            else:
                variants = [marker['word'], f"новый_{marker['word']}", "вариант", ""]
                combo.addItems(variants)
            layout.addWidget(label)
            layout.addWidget(combo)
            item = QListWidgetItem(self.markers_list)
            item.setSizeHint(widget.sizeHint())
            self.markers_list.addItem(item)
            self.markers_list.setItemWidget(item, widget)
            self.variant_widgets.append(combo)

    def update_file_name_options(self):
        self.file_name_combo.clear()
        self.file_name_combo.addItem("Имя шаблона")
        if self.dataframe is not None:
            for display_name in self.column_mapping.values():
                self.file_name_combo.addItem(display_name)

    def _get_selected_columns(self):
        selected_display = [cb.currentText() for cb in self.variant_widgets]
        real_cols = []
        for display in selected_display:
            real = None
            for key, val in self.column_mapping.items():
                if val == display:
                    real = key
                    break
            real_cols.append(real)
        return real_cols

    def generate_filled_document(self, row_index=0, save_all=False):
        if not self.markers:
            QMessageBox.information(self, "Информация", "Маркеры не найдены")
            return None
        if self.dataframe is None or len(self.dataframe) == 0:
            QMessageBox.warning(self, "Ошибка", "Нет данных для заполнения")
            return None
        real_cols = self._get_selected_columns()
        if not save_all:
            if row_index >= len(self.dataframe):
                row_index = 0
            row = self.dataframe.iloc[row_index]
            new_text = self.raw_text
            for idx, marker in enumerate(reversed(self.markers)):
                col = real_cols[len(self.markers) - 1 - idx]
                value = str(row[col]) if col is not None and col in row and not pandas.isna(row[col]) else ""
                new_text = new_text[:marker['start']] + value + new_text[marker['end']:]
            self.generated_text = new_text
            return new_text
        else:
            results = []
            for idx in range(len(self.dataframe)):
                row = self.dataframe.iloc[idx]
                new_text = self.raw_text
                for i, marker in enumerate(reversed(self.markers)):
                    col = real_cols[len(self.markers) - 1 - i]
                    value = str(row[col]) if col is not None and col in row and not pandas.isna(row[col]) else ""
                    new_text = new_text[:marker['start']] + value + new_text[marker['end']:]
                results.append(new_text)
            return results

    def show_example(self):
        if self.dataframe is None or len(self.dataframe) == 0:
            QMessageBox.warning(self, "Ошибка", "Нет данных для примера")
            return
        filled = self.generate_filled_document(row_index=0, save_all=False)
        if filled is not None:
            self.text_edit.setPlainText(filled)

    def generate_all(self):
        if self.dataframe is None or len(self.dataframe) == 0:
            QMessageBox.warning(self, "Ошибка", "Нет данных для генерации")
            return
        from PySide6.QtWidgets import QProgressDialog
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения")
        if not folder:
            return
        folder_path = Path(folder) / self.path.stem
        folder_path.mkdir(exist_ok=True)
        results = self.generate_filled_document(save_all=True)
        if not results:
            return
        name_option = self.file_name_combo.currentText()
        use_template_name = (name_option == "Имя шаблона")
        col_for_name = None
        if not use_template_name:
            for real, display in self.column_mapping.items():
                if display == name_option:
                    col_for_name = real
                    break
        progress = QProgressDialog("Генерация файлов...", "Отмена", 0, len(results), self)
        progress.setWindowModality(Qt.WindowModal)
        for i, text in enumerate(results):
            if progress.wasCanceled():
                break
            progress.setValue(i)
            if use_template_name:
                base_name = f"{self.path.stem}_{i+1}"
            else:
                row = self.dataframe.iloc[i]
                val = row[col_for_name] if col_for_name in row and not pandas.isna(row[col_for_name]) else ""
                base_name = str(val).strip()
                if not base_name:
                    base_name = f"{self.path.stem}_{i+1}"
                base_name = re.sub(r'[\\/*?:"<>|]', "_", base_name)
            filename = f"{base_name}.txt"
            filepath = folder_path / filename
            counter = 1
            while filepath.exists():
                filename = f"{base_name}_{counter}.txt"
                filepath = folder_path / filename
                counter += 1
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(text)
        progress.setValue(len(results))
        QMessageBox.information(self, "Готово", f"Создано {len(results)} файлов в {folder_path}")

class FillWindowFile(QMainWindow):
    def __init__(self, file_path, template_path, parent=None):
        super().__init__(parent)
        self.file_path = file_path
        self.template_path = template_path
        self.dataframe = None
        self.column_mapping = {}
        self.setWindowTitle(f"Выбор данных из: {file_path.name}")
        self.table_widget = None
        self._updating = False
        self._accepting = False
        self.template_win = None
        self.empty_viewer = None
        self.resize(1000, 700)
        self.setup_ui()
        self.load_file()

    def setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)

        btn_layout = QHBoxLayout()
        self.accept_btn = QPushButton("Принять")
        self.add_index_btn = QPushButton("Добавить нумерацию")
        self.add_row_btn = QPushButton("Добавить строку")
        btn_layout.addWidget(self.accept_btn)
        btn_layout.addWidget(self.add_index_btn)
        btn_layout.addWidget(self.add_row_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        splitter = QSplitter(Qt.Horizontal)
        layout.addWidget(splitter)

        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.addWidget(QLabel("Данные:"))
        self.stacked_widget = QWidget()
        self.stacked_layout = QVBoxLayout(self.stacked_widget)
        left_layout.addWidget(self.stacked_widget)
        splitter.addWidget(left_widget)

        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("Теги (колонки):"))
        self.tags_list = QListWidget()
        self.tags_list.setEditTriggers(
            QAbstractItemView.DoubleClicked | QAbstractItemView.EditKeyPressed
        )
        self.tags_list.itemChanged.connect(self.on_tag_renamed)
        right_layout.addWidget(self.tags_list)

        tag_btn_layout = QHBoxLayout()
        self.add_tag_btn = QPushButton("Добавить тег")
        self.remove_tag_btn = QPushButton("Удалить тег")
        tag_btn_layout.addWidget(self.add_tag_btn)
        tag_btn_layout.addWidget(self.remove_tag_btn)
        right_layout.addLayout(tag_btn_layout)

        splitter.addWidget(right_widget)
        splitter.setSizes([700, 300])

        self.accept_btn.clicked.connect(self.accept_selection)
        self.add_tag_btn.clicked.connect(self.add_tag)
        self.remove_tag_btn.clicked.connect(self.remove_tag)
        self.add_index_btn.clicked.connect(self.add_index_column)
        self.add_row_btn.clicked.connect(self.add_row)

    def clear_stacked(self):
        while self.stacked_layout.count():
            child = self.stacked_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

    def load_file(self):
        suffix = self.file_path.suffix.lower()
        if suffix in ['.xlsx', '.xls']:
            self.load_excel()
        elif suffix == '.csv':
            self.load_csv()
        else:
            self.load_text()

    def load_excel(self):
        try:
            self.dataframe = pandas.read_excel(
                self.file_path,
                engine='openpyxl' if self.file_path.suffix == '.xlsx' else 'xlrd',dtype=str
            )
            self.column_mapping = {col: col for col in self.dataframe.columns}
            self.display_dataframe(self.dataframe)
            self.display_tags()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить Excel:\n{e}")

    def load_csv(self):
        try:
            self.dataframe = pandas.read_csv(self.file_path,dtype=str)
            self.column_mapping = {col: col for col in self.dataframe.columns}
            self.display_dataframe(self.dataframe)
            self.display_tags()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить CSV:\n{e}")

    def load_text(self):
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                self.text_content = f.read()
        except UnicodeDecodeError:
            try:
                with open(self.file_path, 'r', encoding='cp1251') as f:
                    self.text_content = f.read()
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось прочитать текст: {e}")
                return
        text_edit = QTextEdit()
        text_edit.setPlainText(self.text_content)
        self.clear_stacked()
        self.stacked_layout.addWidget(text_edit)
        self.text_edit = text_edit
        self.column_mapping = {}
        self.tags_list.clear()
        self.tags_list.addItem("(текстовый файл, теги не определены)")

    def display_dataframe(self, df):
        self._updating = True
        table = QTableWidget()
        table.setRowCount(df.shape[0])
        table.setColumnCount(df.shape[1])
        table.setHorizontalHeaderLabels(df.columns.astype(str))
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                value = df.iat[i, j]
                item = QTableWidgetItem(str(value) if not pandas.isna(value) else "")
                table.setItem(i, j, item)
        table.setEditTriggers(
            QAbstractItemView.DoubleClicked |
            QAbstractItemView.EditKeyPressed |
            QAbstractItemView.AnyKeyPressed
        )
        table.installEventFilter(self)
        table.itemChanged.connect(self.on_cell_changed)
        self.table_widget = table
        self.clear_stacked()
        self.stacked_layout.addWidget(table)
        self._updating = False

    def add_row(self):
        if self.table_widget is None:
            return
        self.add_row_to_table()

    def eventFilter(self, obj, event):
        if obj == self.table_widget and event.type() == QEvent.KeyPress:
            if event.modifiers() == Qt.ControlModifier and event.key() == Qt.Key_V:
                self.paste_from_clipboard()
                return True
        return super().eventFilter(obj, event)

    def update_dataframe_from_table(self):
        for i in range(self.table_widget.rowCount()):
            for j in range(self.table_widget.columnCount()):
                item = self.table_widget.item(i, j)
                if item is not None:
                    val = item.text()
                    self.dataframe.iat[i, j] = val

    def _insert_data_parts(self, rows_data, start_row, start_col):
        if self.table_widget is None:
            return
        self.table_widget.setUpdatesEnabled(False)
        self._updating = True
        for i, row_parts in enumerate(rows_data):
            target_row = start_row + i
            while target_row >= self.table_widget.rowCount():
                self.add_row_to_table()
            for j, cell_text in enumerate(row_parts):
                target_col = start_col + j
                while target_col >= self.table_widget.columnCount():
                    self.add_column_to_table()
                item = self.table_widget.item(target_row, target_col)
                if item is None:
                    item = QTableWidgetItem()
                    self.table_widget.setItem(target_row, target_col, item)
                item.setText(cell_text)
        self.update_dataframe_from_table()
        self.table_widget.setUpdatesEnabled(True)
        self._updating = False

    def paste_from_clipboard(self):
        clipboard = QApplication.clipboard()
        text = clipboard.text()
        if not text:
            return
        if '\n' not in text and '\r' not in text:
            self.paste_single_cell(text)
            return
        lines = re.split(r'\r?\n', text)
        rows_data = []
        for line in lines:
            sep = None
            for s in ['\t', ',', ';', ' ']:
                if s in line:
                    sep = s
                    break
            if sep is None:
                rows_data.append([line])
            else:
                if sep == ' ':
                    parts = [p for p in line.split(' ') if p]
                else:
                    parts = line.split(sep)
                rows_data.append(parts)
        current_row = self.table_widget.currentRow()
        current_col = self.table_widget.currentColumn()
        if current_row < 0 or current_col < 0:
            current_row = 0
            current_col = 0
        self._insert_data_parts(rows_data, current_row, current_col)

    def paste_single_cell(self, text):
        current_row = self.table_widget.currentRow()
        current_col = self.table_widget.currentColumn()
        if current_row < 0 or current_col < 0:
            return
        separators = ['\t', ',', ';', ' ']
        delimiter = None
        for sep in separators:
            if sep in text:
                delimiter = sep
                break
        if delimiter is not None:
            if delimiter == ' ':
                parts = [p for p in text.split(' ') if p]
            else:
                parts = text.split(delimiter)
            rows_data = [[part.strip()] for part in parts]
        else:
            rows_data = [[text]]
        self._insert_data_parts(rows_data, current_row, current_col)

    def on_cell_changed(self, item):
        if self._updating:
            return
        row = item.row()
        col = item.column()
        if row < len(self.dataframe) and col < len(self.dataframe.columns):
            val = item.text()
            if val == "":
                val = pandas.NA
            self.dataframe.iat[row, col] = val

    def add_column_to_table(self):
        self._updating = True
        col_pos = self.table_widget.columnCount()
        self.table_widget.insertColumn(col_pos)
        base_name = "Новая колонка"
        new_col_name = base_name
        counter = 1
        while new_col_name in self.dataframe.columns:
            new_col_name = f"{base_name}_{counter}"
            counter += 1
        self.dataframe[new_col_name] = pandas.NA
        self.column_mapping[new_col_name] = new_col_name
        self.table_widget.setHorizontalHeaderLabels(self.dataframe.columns.astype(str))
        self.display_tags()
        self._updating = False

    def display_tags(self):
        self.tags_list.clear()
        if self.dataframe is not None:
            for col in self.dataframe.columns:
                display_name = self.column_mapping.get(col, col)
                item = QListWidgetItem(display_name)
                item.setData(Qt.UserRole, col)
                self.tags_list.addItem(item)
        else:
            self.tags_list.addItem("(нет данных)")

    def add_index_column(self):
        if self.dataframe is None:
            QMessageBox.warning(self, "Ошибка", "Нет загруженных данных.")
            return
        self._updating = True
        base_name = "Нумерация"
        index_col_name = base_name
        counter = 1
        while index_col_name in self.dataframe.columns:
            index_col_name = f"{base_name}_{counter}"
            counter += 1
        self.dataframe[index_col_name] = [str(i+1) for i in range(len(self.dataframe))]
        self.column_mapping[index_col_name] = index_col_name
        self.display_dataframe(self.dataframe)
        self.display_tags()
        self._updating = False

    def add_row_to_table(self):
        self._updating = True
        row_pos = self.table_widget.rowCount()
        self.table_widget.insertRow(row_pos)
        new_row = {col: pandas.NA for col in self.dataframe.columns}
        self.dataframe = pandas.concat([self.dataframe, pandas.DataFrame([new_row])], ignore_index=True)
        self.table_widget.scrollToItem(self.table_widget.item(row_pos, 0))
        self.table_widget.selectRow(row_pos)
        self._updating = False

    def on_tag_renamed(self, item):
        new_name = item.text().strip()
        if not new_name:
            old_display = self.column_mapping.get(item.data(Qt.UserRole), "")
            item.setText(old_display)
            return
        original_col = item.data(Qt.UserRole)
        if original_col is not None and original_col in self.dataframe.columns:
            if original_col != new_name:
                if new_name in self.dataframe.columns:
                    QMessageBox.warning(self, "Ошибка", f"Колонка с именем '{new_name}' уже существует.")
                    item.setText(self.column_mapping.get(original_col, original_col))
                    return
                self.dataframe.rename(columns={original_col: new_name}, inplace=True)
                self.column_mapping[new_name] = self.column_mapping.pop(original_col)
                item.setData(Qt.UserRole, new_name)
                self.display_dataframe(self.dataframe)
                self.display_tags()

    def add_tag(self):
        if self.dataframe is None:
            QMessageBox.warning(self, "Ошибка", "Нет загруженных данных.")
            return
        self.add_column_to_table()
        for i in range(self.tags_list.count()):
            item = self.tags_list.item(i)
            if i == self.tags_list.count() - 1:
                self.tags_list.setCurrentItem(item)
                self.tags_list.editItem(item)
                break

    def remove_tag(self):
        current_row = self.tags_list.currentRow()
        if current_row < 0:
            QMessageBox.information(self, "Информация", "Выберите тег для удаления.")
            return
        item = self.tags_list.takeItem(current_row)
        if item is None:
            return
        col_name = item.data(Qt.UserRole)
        if col_name is None or col_name not in self.dataframe.columns:
            return
        reply = QMessageBox.question(
            self, "Подтверждение",
            f"Удалить колонку '{col_name}' и тег?",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.No:
            self.display_tags()
            return
        self.dataframe.drop(columns=[col_name], inplace=True)
        if col_name in self.column_mapping:
            del self.column_mapping[col_name]
        self.display_dataframe(self.dataframe)
        self.display_tags()

    def accept_selection(self):
        result = self._show_empty_rows_dialog(
            title="Предупреждение",
            question="В таблице есть пустые значения. Продолжить с ними?",
            continue_label="Продолжить"
        )
        if result == 'cancel' or result == 'show':
            return
        self._accepting = True
        self.template_win = FillWindowTemplate(
            self.template_path,
            dataframe=self.dataframe,
            column_mapping=self.column_mapping
        )
        self.template_win.show()
        self.close()

    def _show_empty_rows_dialog(self, title, question, continue_label):
        empty_rows = self.get_empty_rows()
        if not empty_rows:
            return 'continue'
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle(title)
        msg_box.setText(
            f"{question}\n\n"
            f"Найдено {len(empty_rows)} строк с пустыми значениями."
        )
        continue_btn = msg_box.addButton(continue_label, QMessageBox.AcceptRole)
        show_btn = msg_box.addButton("Показать строки", QMessageBox.ActionRole)
        cancel_btn = msg_box.addButton("Отмена", QMessageBox.RejectRole)
        msg_box.setDefaultButton(cancel_btn)
        msg_box.exec()
        clicked = msg_box.clickedButton()
        if clicked == cancel_btn:
            return 'cancel'
        if clicked == show_btn:
            self.show_empty_rows(empty_rows)
            return 'show'
        return 'continue'

    def get_empty_rows(self):
        empty_rows = []
        if self.dataframe is None:
            return empty_rows
        for i in range(len(self.dataframe)):
            row = self.dataframe.iloc[i]
            if row.isna().any() or (row == "").any():
                empty_rows.append(i)
        return empty_rows

    def show_empty_rows(self, empty_indices):
        if not empty_indices:
            return
        text = f"Строки с пустыми значениями (всего {len(empty_indices)}):\n\n"
        for idx in empty_indices:
            row_data = self.dataframe.iloc[idx].to_dict()
            row_str = f"Строка {idx + 1}: " + ", ".join(f"{k}={v}" for k, v in row_data.items())
            text += row_str + "\n"
        viewer = QTextEdit()
        viewer.setWindowTitle("Строки с пустыми значениями")
        viewer.setPlainText(text)
        viewer.resize(600, 400)
        viewer.show()

    def closeEvent(self, event):
        if self._accepting:
            event.accept()
            return
        result = self._show_empty_rows_dialog(
            title="Предупреждение",
            question="Закрыть окно с пустыми значениями?",
            continue_label="Закрыть"
        )
        if result == 'cancel':
            event.ignore()
            return
        if result == 'show':
            event.ignore()
            return
        event.accept()

def main():
    app = QApplication(sys.argv)
    ui_path = Path("filework.ui")
    window = QUiLoader().load(ui_path)
    window.show()

    dragdrop = DragnDrop(window)

    def on_fill():
        if dragdrop.current_template_path is None:
            QMessageBox.warning(window, "Предупреждение", "Сначала загрузите шаблон!")
            return
        if dragdrop.current_file_path is None:
            QMessageBox.warning(window, "Предупреждение", "Сначала загрузите файл данных!")
            return
        window.hide()
        window.fill_file_window = FillWindowFile(
            dragdrop.current_file_path,
            dragdrop.current_template_path
        )
        window.fill_file_window.show()

    window.Fill.triggered.connect(on_fill)
    sys.exit(app.exec())

if __name__ == "__main__":
    main()